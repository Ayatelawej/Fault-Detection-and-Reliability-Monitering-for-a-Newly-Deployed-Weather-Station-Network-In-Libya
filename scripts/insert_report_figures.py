from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "docs" / "report_drafts"
SOURCE_DOCX = REPORT_DIR / "Ayat_Elawej_EC499_Report_Draft.docx"
OUTPUT_DOCX = REPORT_DIR / "Ayat_Elawej_EC499_Report_Draft_with_figures.docx"
FIGURES_DIR = ROOT / "outputs" / "figures"


PLACEHOLDER_FIGURES = {
    "[Insert map or station distribution figure here": (
        FIGURES_DIR / "report_station_network_map.png",
        6.3,
    ),
    "[Insert system architecture diagram here": (
        FIGURES_DIR / "report_hybrid_fault_detection_architecture.png",
        6.4,
    ),
    "[Insert row-state classifier flow diagram here": (
        FIGURES_DIR / "report_row_state_classification_flow.png",
        6.4,
    ),
    "[Insert event-construction diagram here": (
        FIGURES_DIR / "report_outage_event_construction_flow.png",
        6.4,
    ),
    "[Insert outputs/figures/station_coverage_timeline.png here": (
        FIGURES_DIR / "station_coverage_timeline.png",
        6.4,
    ),
    "[Insert outputs/figures/missingness_heatmap.png here": (
        FIGURES_DIR / "missingness_heatmap.png",
        6.4,
    ),
    "[Insert outputs/figures/station_uptime_bar.png here": (
        FIGURES_DIR / "station_uptime_bar.png",
        6.4,
    ),
    "[Insert outputs/figures/network_offline_fraction_timeline.png here": (
        FIGURES_DIR / "network_offline_fraction_timeline.png",
        6.4,
    ),
    "[Insert example panel here": (
        FIGURES_DIR / "report_itripo33_stuck_wind_panel.png",
        6.4,
    ),
}


def _clear_paragraph(paragraph) -> None:
    for run in paragraph.runs:
        run.text = ""


def _iter_paragraphs(container):
    for paragraph in container.paragraphs:
        yield paragraph
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_paragraphs(cell)


def insert_figures() -> int:
    document = Document(SOURCE_DOCX)
    inserted = 0

    for paragraph in _iter_paragraphs(document):
        text = paragraph.text
        for marker, (figure_path, width_inches) in PLACEHOLDER_FIGURES.items():
            if marker in text:
                if not figure_path.is_file():
                    raise FileNotFoundError(figure_path)
                _clear_paragraph(paragraph)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.add_run().add_picture(
                    str(figure_path),
                    width=Inches(width_inches),
                )
                inserted += 1
                break

    document.save(OUTPUT_DOCX)
    return inserted


def main() -> None:
    inserted = insert_figures()
    print(f"Inserted {inserted} figures into {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
